from docx import Document
import re

# Load the original Word document
doc_path = "docs/contract_template.docx"
doc = Document(doc_path)

# Regex pattern to match single curly brace placeholders like {Placeholder}
pattern = re.compile(r"\{(\w+?)\}")

# Function to replace single brace with double braces
def replace_placeholders(paragraph):
    for run in paragraph.runs:
        run.text = pattern.sub(r"{{ \1 }}", run.text)

# Apply replacement to all paragraphs and table cells
for para in doc.paragraphs:
    replace_placeholders(para)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_placeholders(para)

# Save the modified document
output_path = "templates/contract/contract_template.docx"
doc.save(output_path)

output_path
